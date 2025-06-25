import base64
from faker import Faker
from docx import Document
from docx.shared import Inches
import requests

def encode_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def add_api_response_with_bold_sections(doc, response_content):
    bold_sections = [
        "Accident Description:",
        "Potential Causes:",
        "Immediate Observations:",
        "Recommended Actions:"
    ]
    lines = response_content.split('\n')
    for line in lines:
        paragraph = doc.add_paragraph()
        if any(line.strip().startswith(section) for section in bold_sections):
            run = paragraph.add_run(line)
            run.bold = True
        else:
            paragraph.add_run(line)

def add_text_with_bold_label(doc, label, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(text)

def add_static_sections(doc):
    """Adds static sections like 'Additional Notes' to the document."""
    additional_notes_content = (
        "Send the report to Najm Company for a review of the surveillance footage before and after "
        "the accident by Najm's crash investigator, and forward it to the insurance company, which is "
        "responsible for the claims process."
    )
    add_text_with_bold_label(doc, "Additional Notes: ", additional_notes_content)


def report_generate(image):
    # Initialize Faker
    fake = Faker()

    # API key and image path setup
    api_key = ""
    image_path = image

    # Encode the image for the API request
    base64_image = encode_image(image_path)

    # Headers for the API call
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    # Adjust the payload for the API call
    payload = {
        "model": "gpt-4-vision-preview",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "output must begin with Based on the image, can you provide an answer for all the following :\n-Accident Description:\n- Potential Causes:\n- Immediate Observations:\n- Recommended Actions:"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
            }
        ],
        "max_tokens": 500
    }

    # Make the API request
    response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload)
    response_json = response.json()

    # Extract the relevant response part
    try:
        response_content = response_json['choices'][0]['message']['content']
    except (KeyError, IndexError):
        response_content = "Response content could not be extracted."

    # Create a new DOCX document
    doc = Document()
    doc.add_heading('Image Analysis and API Response', level=1)

    # Add the image to the document
    doc.add_paragraph('Image:')
    doc.add_picture(image_path, width=Inches(6))

    # Generate simulated Time, Date, and Location using Faker
    simulated_time = fake.time()
    simulated_date = fake.date()
    simulated_location = fake.city()

    # Add fake generated details to the document with bold labels
    add_text_with_bold_label(doc, "Time: ", simulated_time)
    add_text_with_bold_label(doc, "Date: ", simulated_date)
    add_text_with_bold_label(doc, "Location: ", simulated_location)

    # Add the API response to the document with specific sections in bold
    add_api_response_with_bold_sections(doc, response_content)

    # Add static sections like 'Additional Notes'
    add_static_sections(doc)

    # Specify the output path and save the document
    output_path = 'Report.docx'
    doc.save(output_path)

    print(f"The document has been saved to: {output_path}")
    return output_path